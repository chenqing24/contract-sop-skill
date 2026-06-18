#!/usr/bin/env python3
"""
合同 MD → docx 模板填充脚本

用法:
    python3 md_to_docx.py <模板.docx> <合同文本_Markdown版.md> <输出.docx>

工作原理:
    1. 读取模板 docx 作为格式骨架
    2. 按段落映射表逐段替换文本 / 删除段落 / 更新表格
    3. 输出新 docx，模板文件保持不变
"""

import sys
import re
from docx import Document


# ═══════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════

def set_para_text(para, text):
    """替换段落全部文字，保留首 run 格式"""
    if not para.runs:
        return
    for run in para.runs:
        run.text = ''
    para.runs[0].text = text
    for run in para.runs[1:]:
        run._element.getparent().remove(run._element)


def del_para(para):
    """从 XML 树彻底删除段落（不是清空文字）"""
    parent = para._element.getparent()
    if parent is not None:
        parent.remove(para._element)


def del_table_row(row):
    """从 XML 树彻底删除表格行"""
    tbl = row._element.getparent()
    if tbl is not None:
        tbl.remove(row._element)


def set_cell_text(cell, text):
    """设置表格单元格文字"""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = text
    elif cell.paragraphs:
        cell.paragraphs[0].text = text


def read_md_paragraphs(md_path):
    """读取 MD 文件，返回段落列表（按 --- 或 ## 分隔）"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    return content


# ═══════════════════════════════════════════════════════
# MAIN — 用户需根据实际模板定制以下映射
# ═══════════════════════════════════════════════════════

if __name__ == '__main__':
    if len(sys.argv) < 4:
        print("用法: python3 md_to_docx.py <模板.docx> <内容.md> <输出.docx>")
        sys.exit(1)

    tpl_path = sys.argv[1]
    md_path = sys.argv[2]
    out_path = sys.argv[3]

    doc = Document(tpl_path)

    # ───────────────────────────────────────────────
    # PHASE 1: 文本替换
    # 格式: set_para_text(doc.paragraphs[索引], "新文本")
    #
    # 替换前先运行以下代码获取模板的段落映射:
    #   for i, p in enumerate(doc.paragraphs):
    #       if p.text.strip():
    #           print(f"P{i:03d}: {p.text[:60]}")
    # ───────────────────────────────────────────────

    # 示例:
    # set_para_text(doc.paragraphs[28], "乙方为甲方提供以下内容：...")

    # ───────────────────────────────────────────────
    # PHASE 2: 删除不需要的段落
    # 必须倒序删除！
    #
    # DELETE_INDICES = [索引1, 索引2, ...]
    # for idx in sorted(DELETE_INDICES, reverse=True):
    #     if idx < len(doc.paragraphs):
    #         del_para(doc.paragraphs[idx])
    #         print(f"  🗑  删除 P{idx:03d}")
    # ───────────────────────────────────────────────

    # ───────────────────────────────────────────────
    # PHASE 3: 表格更新
    #
    # table = doc.tables[N]
    # set_cell_text(row.cells[M], "新内容")
    #
    # 删除多余行:
    # for row_idx in sorted(rows_to_delete, reverse=True):
    #     del_table_row(table.rows[row_idx])
    # ───────────────────────────────────────────────

    # ───────────────────────────────────────────────
    # 保存输出
    # ───────────────────────────────────────────────
    doc.save(out_path)
    print(f"✅ 合同已生成: {out_path}")

    # 最终验证
    print(f"   段落数: {len(doc.paragraphs)}")
    print(f"   表格数: {len(doc.tables)}")
